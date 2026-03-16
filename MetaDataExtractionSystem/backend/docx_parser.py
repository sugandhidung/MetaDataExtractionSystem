"""
DOCX text extraction module.
Extracts full text content from .docx files using python-docx.
"""

from pathlib import Path
from docx import Document


def extract_text_from_docx(file_path: str | Path) -> str:
    """
    Extract all text from a .docx file, including paragraphs and tables.
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        Full text content of the document.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(str(file_path))
    text_parts = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)
